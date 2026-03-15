"""Manuscript ingest for Quillan2 — 5E: Import/Ingest Existing Manuscripts.

Parses an existing Markdown, plain-text, or DOCX manuscript into chapters and
beats, writes Beat_Draft.md files, generates a stub Outline.yaml and
dependency_map.json, and optionally triggers the full planning pipeline
(Story_Spine, Character_Arcs, etc.).

Supported formats:
  .md / .txt  — headings (# or ##) delimit chapters
  .docx       — Heading 1/Heading 2 paragraph styles delimit chapters

Usage:
    from quillan.ingest import ingest_manuscript
    story_name = await ingest_manuscript(Path("my_book.md"), paths, ...)
"""

from __future__ import annotations

import json
import re
from collections.abc import Callable
from pathlib import Path
from typing import TYPE_CHECKING

import yaml

if TYPE_CHECKING:
    from quillan.config import Settings
    from quillan.llm import LLMClient
    from quillan.paths import Paths

_HEADING_RE = re.compile(r"^#{1,2}\s+(.+)$", re.MULTILINE)

# Style sample extraction: first N chapters, capped so samples.md stays readable
_STYLE_SAMPLE_CHAPTERS = 3
_STYLE_SAMPLE_MAX_CHARS = 6000


# ── Text parsing ──────────────────────────────────────────────────────────────


def parse_markdown(text: str) -> list[dict]:
    """Split Markdown/plain text into chapters by level-1/2 headings.

    Returns a list of ``{'title': str, 'text': str}`` dicts.
    Text before the first heading is kept as a 'Preface' chapter if non-empty.
    If no headings are found the entire text is returned as a single chapter.
    """
    text = text.replace("\r\n", "\n")
    splits: list[tuple[int, str]] = []  # (match_start, heading_text)
    for m in _HEADING_RE.finditer(text):
        splits.append((m.start(), m.group(1).strip()))

    if not splits:
        stripped = text.strip()
        return [{"title": "Chapter 1", "text": stripped}] if stripped else []

    chapters: list[dict] = []

    # Text before the first heading
    pre = text[: splits[0][0]].strip()
    if pre:
        chapters.append({"title": "Preface", "text": pre})

    for i, (start, title) in enumerate(splits):
        end = splits[i + 1][0] if i + 1 < len(splits) else len(text)
        # Skip the heading line itself
        newline_pos = text.find("\n", start)
        heading_end = newline_pos + 1 if newline_pos != -1 and newline_pos < end else end
        body = text[heading_end:end].strip()
        if body:
            chapters.append({"title": title, "text": body})

    return chapters


def parse_docx(path: Path) -> list[dict]:
    """Parse a DOCX file into chapters using Heading 1/Heading 2 styles.

    Raises ImportError if python-docx is not installed.
    """
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX import. "
            "Install it: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    chapters: list[dict] = []
    current_title: str | None = None
    current_paras: list[str] = []

    _HEADING_STYLES = {"heading 1", "heading 2", "heading1", "heading2"}

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower().strip()
        text = para.text.strip()
        if not text:
            continue
        if style_name in _HEADING_STYLES:
            if current_paras:
                chapters.append(
                    {
                        "title": current_title or "Chapter",
                        "text": "\n\n".join(current_paras),
                    }
                )
                current_paras = []
            current_title = text
        else:
            current_paras.append(text)

    if current_paras:
        chapters.append(
            {
                "title": current_title or "Chapter 1",
                "text": "\n\n".join(current_paras),
            }
        )

    if not chapters:
        # No headings found — treat whole doc as one chapter
        all_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if all_text:
            chapters.append({"title": "Chapter 1", "text": all_text})

    return chapters


# ── Beat clustering ───────────────────────────────────────────────────────────


def cluster_into_beats(chapter_text: str, target_words: int = 1500) -> list[str]:
    """Split chapter text into beat-sized chunks at paragraph boundaries.

    Tries to stay close to *target_words* per beat.  Never splits within a
    paragraph — a paragraph that exceeds *target_words* on its own becomes
    its own beat.
    """
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", chapter_text) if p.strip()]
    if not paragraphs:
        return []

    beats: list[str] = []
    current_paras: list[str] = []
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())
        if current_words + para_words > target_words and current_paras:
            beats.append("\n\n".join(current_paras))
            current_paras = [para]
            current_words = para_words
        else:
            current_paras.append(para)
            current_words += para_words

    if current_paras:
        beats.append("\n\n".join(current_paras))

    return beats


# ── Style sample extraction ───────────────────────────────────────────────────


def _build_style_samples(chapters_with_beats: list[dict]) -> str:
    """Concatenate the first _STYLE_SAMPLE_CHAPTERS chapters into a samples string.

    Chapters are separated by '---' dividers. The result is capped at
    _STYLE_SAMPLE_MAX_CHARS so samples.md stays manageable.
    Returns an empty string if no chapters are available.
    """
    selected = chapters_with_beats[:_STYLE_SAMPLE_CHAPTERS]
    if not selected:
        return ""

    parts: list[str] = []
    for ch in selected:
        title = ch["title"]
        text = "\n\n".join(ch["beats"]).strip()
        if text:
            parts.append(f"### {title}\n\n{text}")

    combined = "\n\n---\n\n".join(parts)
    if len(combined) > _STYLE_SAMPLE_MAX_CHARS:
        combined = combined[:_STYLE_SAMPLE_MAX_CHARS] + "\n...(truncated)"
    return combined


# ── Outline / spec / dep-map builders ────────────────────────────────────────


def _build_ingest_outline(chapters_with_beats: list[dict], title: str) -> dict:
    """Produce an Outline.yaml-compatible dict from ingest chapter/beat data."""
    chapters_yaml = []
    for ch_num, ch in enumerate(chapters_with_beats, 1):
        beats_yaml = []
        for b_num, beat_text in enumerate(ch["beats"], 1):
            word_count = len(beat_text.split())
            beats_yaml.append(
                {
                    "beat_id": f"C{ch_num}-S1-B{b_num}",
                    "title": f"Beat {b_num}",
                    "goal": "(imported beat — edit as needed)",
                    "setting": "(imported)",
                    "characters": [],
                    "word_count_target": max(500, round(word_count / 100) * 100),
                }
            )
        chapters_yaml.append(
            {
                "chapter": ch_num,
                "title": ch["title"],
                "beats": beats_yaml,
            }
        )
    return {
        "title": title,
        "genre": "Fiction",
        "theme": "(imported — edit as needed)",
        "source": "imported",
        "chapters": chapters_yaml,
    }


def _stub_beat_spec(beat_id: str, title: str, word_count: int) -> dict:
    return {
        "beat_id": beat_id,
        "title": title,
        "goal": "(imported beat — edit as needed)",
        "setting": "(imported)",
        "characters": [],
        "pov_character": "",
        "tone": "neutral",
        "word_count_target": word_count,
        "emotional_beat": "(imported)",
        "theme_payoff": "(imported)",
        "pacing": "medium",
        "scope": ["(imported)"],
        "out_of_scope": [],
        "rules": [],
        "dependencies": [],
        "arc_position": "",
    }


def _build_dep_map(chapters_with_beats: list[dict]) -> dict:
    """Build a simple linear dependency map: each beat depends on the previous."""
    deps: dict[str, list[str]] = {}
    prev: str | None = None
    for ch_num, ch in enumerate(chapters_with_beats, 1):
        for b_num in range(1, len(ch["beats"]) + 1):
            beat_id = f"C{ch_num}-S1-B{b_num}"
            deps[beat_id] = [prev] if prev else []
            prev = beat_id
    return {"dependencies": deps}


# ── Main ingest entrypoint ────────────────────────────────────────────────────


async def ingest_manuscript(
    source_path: Path,
    paths: "Paths",
    world: str,
    canon: str,
    series: str,
    story: str,
    llm: "LLMClient | None" = None,
    settings: "Settings | None" = None,
    target_words_per_beat: int = 1500,
    run_planning: bool = False,
    on_progress: Callable[[str], None] | None = None,
) -> str:
    """Ingest a manuscript file and create a Quillan2 story from it.

    Writes Beat_Draft.md files for all inferred beats, generates a stub
    Outline.yaml, beat_spec.yaml files, and a linear dependency_map.json.
    Optionally runs the full planning pipeline (Story_Spine, Character_Arcs,
    etc.) if *run_planning* is True and *llm* / *settings* are provided.

    Returns the story name as written to disk.
    """
    from quillan.io import atomic_write

    def _prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)

    source_path = Path(source_path)
    suffix = source_path.suffix.lower()

    _prog(f"Parsing {source_path.name}…")

    if suffix in (".md", ".txt"):
        chapters = parse_markdown(
            source_path.read_text(encoding="utf-8", errors="replace")
        )
    elif suffix == ".docx":
        chapters = parse_docx(source_path)
    else:
        raise ValueError(
            f"Unsupported file type: {suffix!r}. Supported: .md, .txt, .docx"
        )

    if not chapters:
        raise ValueError("No readable content found in manuscript")

    _prog(f"Found {len(chapters)} chapter(s). Clustering into beats…")

    chapters_with_beats: list[dict] = []
    total_beats = 0
    for ch in chapters:
        beats = cluster_into_beats(ch["text"], target_words=target_words_per_beat)
        if not beats:
            continue
        chapters_with_beats.append({"title": ch["title"], "beats": beats})
        total_beats += len(beats)

    if not chapters_with_beats:
        raise ValueError("No content could be extracted from manuscript")

    _prog(f"Writing {total_beats} beat draft(s) to disk…")

    # Initialise story directory tree
    for d in [
        paths.story_input(world, canon, series, story),
        paths.story_planning(world, canon, series, story),
        paths.story_structure(world, canon, series, story),
        paths.story_beats(world, canon, series, story),
        paths.story_state(world, canon, series, story),
        paths.story_export(world, canon, series, story),
        paths.story_continuity(world, canon, series, story),
        paths.queue_dir(world, canon, series, story),
    ]:
        d.mkdir(parents=True, exist_ok=True)

    # Copy source manuscript into story input/
    import shutil
    dest_input = paths.story_input(world, canon, series, story) / source_path.name
    if not dest_input.exists():
        shutil.copy2(source_path, dest_input)

    # Write beat drafts and stub specs
    title = story.replace("_", " ").title()
    for ch_num, ch in enumerate(chapters_with_beats, 1):
        for b_num, beat_text in enumerate(ch["beats"], 1):
            beat_id = f"C{ch_num}-S1-B{b_num}"
            word_count = len(beat_text.split())

            draft_path = paths.beat_draft(world, canon, series, story, beat_id)
            paths.ensure(draft_path)
            atomic_write(draft_path, beat_text)

            spec = _stub_beat_spec(
                beat_id,
                f"Beat {b_num}",
                max(500, round(word_count / 100) * 100),
            )
            spec_path = paths.beat_spec(world, canon, series, story, beat_id)
            paths.ensure(spec_path)
            atomic_write(spec_path, yaml.dump(spec, allow_unicode=True, sort_keys=False))

    # Write Outline.yaml
    _prog("Writing Outline.yaml…")
    outline = _build_ingest_outline(chapters_with_beats, title)
    outline_path = paths.outline(world, canon, series, story)
    paths.ensure(outline_path)
    atomic_write(outline_path, yaml.dump(outline, allow_unicode=True, sort_keys=False))

    # Write dependency_map.json
    _prog("Writing dependency map…")
    dep_map = _build_dep_map(chapters_with_beats)
    dep_path = paths.dependency_map(world, canon, series, story)
    paths.ensure(dep_path)
    atomic_write(dep_path, json.dumps(dep_map, indent=2))

    # ── Auto-extract style samples from first N chapters ─────────────────
    samples_text = _build_style_samples(chapters_with_beats)
    if samples_text:
        from quillan.io import atomic_write as _atomic_write

        samples_path = paths.style_samples(world, canon, series, story)
        paths.ensure(samples_path)
        _atomic_write(samples_path, samples_text)
        _prog(f"Style samples written ({len(samples_text.split())} words from first "
              f"{min(len(chapters_with_beats), _STYLE_SAMPLE_CHAPTERS)} chapter(s)).")

        if llm is not None and settings is not None:
            _prog("Extracting style fingerprint…")
            from quillan.structure.style import extract_style_profile
            profile_path = await extract_style_profile(
                paths, world, canon, series, story, llm, settings
            )
            if profile_path:
                _prog(f"Style profile written: {profile_path.name}")
            else:
                _prog("  Warning: style profile extraction failed (will use raw samples).")

    # Optional: run full planning pipeline
    if run_planning and llm is not None and settings is not None:
        _prog("Generating planning artifacts (Story_Spine, Character_Arcs, …)…")
        import asyncio

        from quillan.structure.world import create_world_if_missing, build_canon_packet
        from quillan.structure.story_spine import generate_story_spine
        from quillan.structure.character_arcs import generate_character_arcs
        from quillan.structure.subplots import generate_subplot_register
        from quillan.structure.conflicts import generate_conflict_map

        await create_world_if_missing(paths, world, llm)

        results = await asyncio.gather(
            generate_story_spine(paths, world, canon, series, story, llm),
            generate_character_arcs(paths, world, canon, series, story, llm),
            generate_subplot_register(paths, world, canon, series, story, llm),
            generate_conflict_map(paths, world, canon, series, story, llm),
            return_exceptions=True,
        )
        for r in results:
            if isinstance(r, Exception):
                _prog(f"  Warning: planning artifact generation failed: {r}")

        try:
            await build_canon_packet(paths, world, canon, series, story, llm)
        except Exception as exc:
            _prog(f"  Warning: Canon Packet generation failed: {exc}")

        _prog("Planning artifacts generated.")

    _prog(f"Ingest complete — story: {story!r}  beats: {total_beats}")
    return story
